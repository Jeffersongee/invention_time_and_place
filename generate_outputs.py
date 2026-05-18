#!/usr/bin/env python3
"""Generate inventions_with_dates.csv and inventions_full_analysis.docx from checkpoint."""
import json
import csv
import sys

# ── Load checkpoint ──────────────────────────────────────────────────────────
with open("analysis_checkpoint.jsonl", "r") as f:
    entries = [json.loads(line) for line in f]

print(f"Loaded {len(entries)} entries")

# ── CSV ──────────────────────────────────────────────────────────────────────
CSV_FIELDS = [
    "invention", "year", "inventor", "category", "category2",
    "earliest_plausible", "earliest_straightforward",
    "actual_location", "possible_locations",
    "geographic_flag", "geographic_note",
    "confidence",
]

with open("inventions_with_dates.csv", "w", newline="", encoding="utf-8") as f:
    writer = csv.DictWriter(f, fieldnames=CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    writer.writerows(entries)

print("Wrote inventions_with_dates.csv")

# ── DOCX ─────────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed; installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("Invention Dating Analysis: Full Results", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    f"Complete analysis of {len(entries)} inventions. "
    "Each entry includes earliest plausible date, earliest straightforward date, "
    "geographic alignment assessment, and full analytical text."
)
doc.add_paragraph()

for i, entry in enumerate(entries):
    # Section heading: index + invention name
    inv_name = entry.get("invention", "Unknown")
    heading = doc.add_heading(f"{i+1}. {inv_name}", level=1)

    # Summary table as paragraphs
    def add_field(label, value):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(str(value) if value else "—")

    add_field("Year", entry.get("year", ""))
    add_field("Inventor", entry.get("inventor", ""))
    add_field("Category", entry.get("category", ""))
    if entry.get("category2"):
        add_field("Category 2", entry.get("category2", ""))
    add_field("Earliest plausible", entry.get("earliest_plausible", ""))
    add_field("Earliest straightforward", entry.get("earliest_straightforward", ""))
    add_field("Confidence", entry.get("confidence", ""))
    add_field("Actual location", entry.get("actual_location", ""))
    add_field("Possible locations", entry.get("possible_locations", ""))
    add_field("Geographic flag", entry.get("geographic_flag", ""))
    add_field("Geographic note", entry.get("geographic_note", ""))

    doc.add_paragraph()

    # Full analysis text
    full_text = entry.get("full_text", "")
    if full_text:
        doc.add_heading("Full Analysis", level=2)
        for line in full_text.strip().split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)

    # Page break after each entry except the last
    if i < len(entries) - 1:
        doc.add_page_break()

doc.save("inventions_full_analysis.docx")
print("Wrote inventions_full_analysis.docx")
